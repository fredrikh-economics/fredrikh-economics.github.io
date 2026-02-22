"""
make_doc.py
Generates ModernDID_Documentation.docx describing the full pipeline,
estimation methods (pros/cons), and setup.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x3A, 0x5F)
DARK   = RGBColor(0x22, 0x22, 0x22)
GREY   = RGBColor(0x55, 0x55, 0x55)
GREEN  = RGBColor(0x1A, 0x6B, 0x3C)
RED    = RGBColor(0x8B, 0x1A, 0x1A)
LIGHT  = RGBColor(0xF2, 0xF5, 0xF9)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, sides=("top","bottom","left","right"),
                    color="1F3A5F", sz="6"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def h1(text):
    p = doc.add_heading(text, level=1)
    run = p.runs[0]
    run.font.color.rgb = NAVY
    run.font.size      = Pt(18)
    run.bold           = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.color.rgb = NAVY
    run.font.size      = Pt(13)
    run.bold           = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    run = p.runs[0]
    run.font.color.rgb = DARK
    run.font.size      = Pt(11)
    run.bold           = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, italic=False, color=DARK):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size       = Pt(10.5)
    run.font.color.rgb  = color
    run.font.italic     = italic
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.space_before = Pt(0)
    return p

def bullet(text, level=0, bold_prefix=None, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + 0.2*level)
    p.paragraph_format.space_after  = Pt(3)
    text_color = color if color else DARK
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold           = True
        run.font.size      = Pt(10.5)
        run.font.color.rgb = DARK
        run2 = p.add_run(" " + text)
        run2.font.size      = Pt(10.5)
        run2.font.color.rgb = text_color
    else:
        run = p.add_run(text)
        run.font.size      = Pt(10.5)
        run.font.color.rgb = text_color
    return p

def code_inline(p, text):
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x5E)
    return run

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("Modern Difference-in-Differences Event-Study Pipeline")
tr.font.size      = Pt(22)
tr.font.bold      = True
tr.font.color.rgb = NAVY

doc.add_paragraph()

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run("Documentation of Methods, Setup, and Implementation")
sr.font.size      = Pt(13)
sr.font.color.rgb = GREY
sr.font.italic    = True

doc.add_paragraph()

ap = doc.add_paragraph()
ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = ap.add_run(
    "Outcomes: lnva_L  ·  lsize  ·  edu_h  ·  exp_sales  ·  imp_sales\n"
    "Window: t = −4 to +9  ·  Reference: t = −1"
)
ar.font.size      = Pt(11)
ar.font.color.rgb = DARK

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

h1("1. Overview")
body(
    "This pipeline estimates the causal effect of foreign acquisitions on "
    "Swedish firms using modern difference-in-differences (DID) methods. "
    "The dataset is a propensity-score-matched (PSM) sample of treated "
    "(acquired) and control (never acquired) firms observed over multiple "
    "calendar years. Five outcome variables are studied: log value added per "
    "worker (lnva_L), firm size (lsize), share of high-educated workers "
    "(edu_h), export sales (exp_sales), and import sales (imp_sales)."
)
body(
    "The pipeline addresses a key challenge in the modern DID literature: "
    "when treatment is staggered across cohorts (different firms are acquired "
    "in different years), the classical two-way fixed-effects (TWFE) estimator "
    "can be biased because it uses already-treated units as implicit controls. "
    "Four complementary estimators are run to check robustness: TWFE, "
    "Sun-Abraham, Callaway-Sant'Anna, and Borusyak-Jaravel-Spiess. "
    "HonestDiD sensitivity analysis is added to assess how large a violation "
    "of parallel trends would be needed to overturn the results."
)


# ══════════════════════════════════════════════════════════════════════════════
# 2. DATA SETUP
# ══════════════════════════════════════════════════════════════════════════════

h1("2. Data Setup and Key Variables")

h2("2.1  Source Data")
body(
    "The starting point is the file "
    "PSM_matched_T_C_firms_all_years_auto_country_info_trade_reg_data.dta. "
    "This is the output of a prior propensity-score matching step that pairs "
    "each acquired firm with one or more never-acquired firms that were "
    "similar in observable characteristics before the acquisition."
)

h2("2.2  Variables Constructed in the Pipeline")

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = tbl.rows[0].cells
hdr[0].text = "Variable"
hdr[1].text = "Definition"
for cell in hdr:
    set_cell_bg(cell, "1F3A5F")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.bold      = True
            run.font.size      = Pt(10)

rows = [
    ("id",            "Firm identifier (renamed from firm)"),
    ("t",             "Calendar year (renamed from yr)"),
    ("treated_firm",  "1 if the firm is ever acquired during the sample; 0 for controls. "
                      "Time-invariant: computed as max(treatment) within firm."),
    ("G",             "Cohort year = the calendar year in which year_st_firm == 0 for "
                      "treated firms. Missing for controls (never-treated)."),
    ("reltime",       "Event time = t − G for treated firms; missing for controls. "
                      "Ranges from −4 to +9 in the analysis window."),
]
for var, defn in rows:
    row = tbl.add_row().cells
    row[0].text = var
    row[1].text = defn
    for para in row[0].paragraphs:
        for run in para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
    set_cell_bg(row[0], "EEF2F7")

doc.add_paragraph()

h2("2.3  Sample Restrictions")
body(
    "Each outcome uses a dedicated sample flag (control_sample, "
    "control_sample_lsize, etc.) defined in the data-creation step. "
    "All estimations additionally restrict to observations within the "
    "event-study window (sample_4_9 == 1) and to calendar years up to 2019."
)

h2("2.4  Sanity Checks")
body(
    "Before estimation a series of seven data-quality checks are run and "
    "per-firm pass/fail flags are created:"
)
checks = [
    ("1. Unique id-t",          "No duplicate firm-year observations."),
    ("2. Treatment coding",     "treatment ∈ {0, 1} and non-missing for all observations."),
    ("3. Controls never treated","Control firms have G and reltime missing throughout."),
    ("4. One event year",       "Each treated firm has exactly one row where year_st_firm = 0."),
    ("5. Unique cohort G",      "Each treated firm has a single, consistent G value."),
    ("6. year_st_firm consistency","year_st_firm = t − G for all treated observations."),
    ("7. Swedish ownership",    "Control firms are always domestically owned (requires fof variable)."),
]
for label, desc in checks:
    bullet(desc, bold_prefix=label)

body(
    "Firms failing any check are saved to bad_ids.dta. Setting DROP_BAD = 1 "
    "removes them and saves a clean dataset. Running with USE_CLEAN = 1 then "
    "applies all estimations to the clean sample."
)


# ══════════════════════════════════════════════════════════════════════════════
# 3. FILE STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

h1("3. File Structure")
body(
    "The pipeline consists of eight Stata do-files, all located in the stata/ "
    "subfolder. Only 00_prep.do needs to be run; it calls the estimation "
    "modules in sequence."
)

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
h2c = tbl2.rows[0].cells
for i, txt in enumerate(["File", "Role", "Key outputs"]):
    h2c[i].text = txt
    set_cell_bg(h2c[i], "1F3A5F")
    for para in h2c[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.bold      = True
            run.font.size      = Pt(10)

file_rows = [
    ("00_prep.do",              "Master + setup + data build",
     "analysis_modern_did_inputs.dta, bad_ids.dta, sets $ANALYSIS_FILE"),
    ("01_twfe.do",              "TWFE event study",
     "TWFE_<y>_-4_9.xlsx / .png  (one per outcome)"),
    ("02_sun_abraham.do",       "Sun-Abraham estimator",
     "SA_<y>_-4_9.xlsx / .png"),
    ("03_csdid.do",             "Callaway-Sant'Anna estimator",
     "CSDID_<y>_-4_9.xlsx / .png"),
    ("04_did_imputation.do",    "BJS imputation estimator",
     "did_imputation_-4_9.log"),
    ("05_extra_tests.do",       "Placebo, Bacon, wild bootstrap",
     "Printed to Stata output / Results window"),
    ("06_comparison_figure.do", "Overlay plot TWFE vs SA vs CSDID",
     "COMPARE_lnva_L_-4_9.png"),
    ("07_honest_did.do",        "HonestDiD sensitivity",
     "HonestDiD_M_<y>.xlsx/.png, HonestDiD_RM_<y>.xlsx/.png"),
    ("08_figures.do",           "Standalone figure generator",
     "Fig_TWFE/SA/CSDID/Compare_<y>.png, Fig_All_TWFE/SA/CSDID.png, Fig_HonestDiD_*.png"),
]
for fname, role, outputs in file_rows:
    r = tbl2.add_row().cells
    r[0].text = fname
    r[1].text = role
    r[2].text = outputs
    for para in r[0].paragraphs:
        for run in para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
    set_cell_bg(r[0], "EEF2F7")

doc.add_paragraph()

h2("3.1  User Settings (top of 00_prep.do)")

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = "Table Grid"
hdr3 = tbl3.rows[0].cells
for i, txt in enumerate(["Global / Toggle", "Default", "Description"]):
    hdr3[i].text = txt
    set_cell_bg(hdr3[i], "1F3A5F")
    for para in hdr3[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.bold      = True
            run.font.size      = Pt(10)

settings = [
    ("$ROOT",               "UNC path",   "Root folder containing data and do-files"),
    ("$DODIR",              "$ROOT\\stata","Folder containing 01_twfe.do … 07_honest_did.do"),
    ("$DROP_BAD",           "0",          "1 = drop bad ids and save clean dataset"),
    ("$USE_CLEAN",          "0",          "1 = use clean dataset (requires DROP_BAD=1 first)"),
    ("$tmin / $tmax",       "-4 / 9",     "Event-study window endpoints"),
    ("$winflag",            "sample_4_9==1","Observation-level window restriction"),
    ("$CTRL_DEFAULT",       "p_firm_age …","Covariate list passed to all estimators"),
    ("$OUTCOMES",           "lnva_L …",   "Space-separated list of outcome variables"),
    ("$SFLAG_<outcome>",    "varies",     "Per-outcome sample flag (e.g. control_sample==1)"),
    ("$HONESTDID_NPRE",     "3",          "Number of pre-treatment dummies (t=−4,−3,−2)"),
    ("$HONESTDID_NPOST",    "10",         "Number of post-treatment dummies (t=0…9)"),
    ("$HONESTDID_MVEC_SMOOTH","0(0.5)2", "Grid for smoothness restriction M"),
    ("$HONESTDID_MVEC_RM",  "0(0.25)1",  "Grid for relative-magnitudes Mbar"),
]
for gname, default, desc in settings:
    r = tbl3.add_row().cells
    r[0].text = gname
    r[1].text = default
    r[2].text = desc
    for para in r[0].paragraphs:
        for run in para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
    set_cell_bg(r[0], "EEF2F7")

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 4. ESTIMATION METHODS
# ══════════════════════════════════════════════════════════════════════════════

h1("4. Estimation Methods")

body(
    "All four estimators share the same event-study design: a balanced window "
    "of t = −4 to +9 relative to acquisition, with t = −1 as the normalisation "
    "period. Pre-treatment coefficients (t = −4, −3, −2) serve as a parallel-trends "
    "test; post-treatment coefficients (t = 0 … +9) trace out the dynamic "
    "treatment effect."
)

# ── 4.1 TWFE ─────────────────────────────────────────────────────────────────
h2("4.1  Two-Way Fixed Effects (TWFE) — 01_twfe.do")

body(
    "The classical panel estimator absorbs firm fixed effects (controlling for "
    "all time-invariant heterogeneity) and year fixed effects (controlling for "
    "aggregate shocks common to all firms). The event-study version adds "
    "binary indicator variables for each combination of treated firm × event time."
)

h3("Model")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
p.add_run("Y").font.italic = True
p.add_run("ᵢₜ = αᵢ + λₜ + ")
bold = p.add_run("∑")
bold.font.bold = True
p.add_run("  βₖ · 𝟙(reltime = k) · treated_firmᵢ + X′ᵢₜ γ + εᵢₜ")
p.paragraph_format.space_after = Pt(6)

body(
    "where αᵢ are firm FEs, λₜ are year FEs, k ∈ {−4,−3,−2, 0,…,9} "
    "(k = −1 omitted as reference), and Xᵢₜ are time-varying controls. "
    "Standard errors are clustered at the firm level. Estimated with reghdfe."
)

h3("Implementation note — explicit binary dummies")
body(
    "Because reltime is missing for control firms, using Stata's factor-variable "
    "interaction syntax (ib(-1).reltime##i.treated_firm) would silently drop all "
    "control observations from the regression. The pipeline instead generates "
    "explicit binary dummies d_m4, d_m3, d_m2, d_p0 … d_p9 that are set to 0 "
    "for control observations. Controls therefore remain in the regression and "
    "contribute to the estimation of the fixed effects."
)

h3("Pros")
for pro in [
    "Simple, transparent, and widely used — referees and readers will be familiar with it.",
    "Computationally fast; reghdfe handles large datasets efficiently.",
    "Easy to include rich covariate controls and arbitrary FE structures.",
    "The joint F-test on pre-treatment dummies is a standard and intuitive pre-trend test.",
]:
    bullet(pro, color=GREEN)

h3("Cons")
for con in [
    "With staggered adoption, the TWFE estimator is a weighted average of all "
    "pairwise 2×2 DID comparisons. Some weights can be negative when treatment "
    "effects are heterogeneous across cohorts (Goodman-Bacon 2021). This means "
    "the estimator can be biased — even in the wrong direction.",
    "Already-treated firms are used as implicit control units for later-treated "
    "cohorts, violating the clean-comparison requirement.",
    "Gives a single pooled post-treatment average rather than cohort-specific effects.",
]:
    bullet(con, color=RED)

# ── 4.2 Sun-Abraham ───────────────────────────────────────────────────────────
h2("4.2  Sun-Abraham (SA) — 02_sun_abraham.do")

body(
    "Sun and Abraham (2021) show that the TWFE event-study estimator is a "
    "weighted average of cohort-specific average treatment effects (CATTs) "
    "with potentially negative weights. Their estimator instead computes "
    "CATTs separately for each cohort-time cell and then aggregates them "
    "using shares of each cohort in the sample — a procedure called "
    "interaction-weighted (IW) estimation."
)

h3("Estimator")
body(
    "The IW estimator runs a fully-saturated regression that interacts each "
    "event-time indicator with cohort dummies. The aggregate event-study "
    "coefficients are then formed as a weighted average of the cohort-specific "
    "coefficients, using cohort shares as weights. Implemented via "
    "eventstudyinteract (Stata)."
)

h3("Pros")
for pro in [
    "Robust to heterogeneous treatment effects across cohorts — the aggregation "
    "weights are always non-negative.",
    "Directly comparable to TWFE in event-study plot format; easy to read.",
    "Provides a formal test of whether TWFE weights are problematic (compare "
    "SA coefficients to TWFE coefficients: large divergence signals heterogeneity).",
    "Never-treated firms are the comparison group only, so no 'forbidden' comparisons.",
]:
    bullet(pro, color=GREEN)

h3("Cons")
for con in [
    "Requires a sufficiently large never-treated group to serve as the clean "
    "comparison; not applicable if all units are eventually treated.",
    "Efficiency loss compared to TWFE: more parameters are estimated "
    "(one per cohort × event-time cell), so standard errors are typically wider.",
    "Pre-trend test has lower power than TWFE because of the larger model.",
    "Cohort-specific estimates can be very noisy for small cohorts.",
]:
    bullet(con, color=RED)

# ── 4.3 CSDID ─────────────────────────────────────────────────────────────────
h2("4.3  Callaway-Sant'Anna (CSDID) — 03_csdid.do")

body(
    "Callaway and Sant'Anna (2021) propose a non-parametric DID estimator "
    "that targets cohort-time average treatment effects ATT(g, t) — the "
    "average effect on cohort g at time t — and then aggregates them "
    "in user-specified ways (by event time, by calendar time, or into a "
    "single overall ATT). The pipeline uses the doubly-robust IPW estimator "
    "(method(dripw)) which combines outcome-regression and propensity-score "
    "approaches; it is consistent if either component is correctly specified."
)

h3("Aggregation")
body(
    "estat event, window(-4 9) produces event-time aggregate ATTs — the "
    "average over cohorts of ATT(g, t−g = k) — which are directly comparable "
    "to the TWFE event-study coefficients."
)

h3("Pros")
for pro in [
    "Formally targets well-defined causal estimands (cohort-time ATTs).",
    "Doubly-robust: consistent if either the propensity score model or the "
    "outcome model is correctly specified.",
    "No negative-weight problem: aggregation is always over non-negative weights.",
    "Allows flexible choice of comparison group: never-treated, not-yet-treated, "
    "or both.",
    "Supports parallel-trends conditional on covariates (conditional DID).",
]:
    bullet(pro, color=GREEN)

h3("Cons")
for con in [
    "Computationally intensive, especially with many cohorts or covariates.",
    "Inference requires a large number of clusters for the cluster bootstrap to "
    "be reliable.",
    "The doubly-robust estimator can behave poorly in finite samples when "
    "propensity scores are near 0 or 1 (extreme overlap violations).",
    "G must be coded as 0 for never-treated (not missing) — a common source "
    "of errors when interfacing with other estimation packages.",
    "The aggregation step (estat event) can produce label-parsing issues "
    "across csdid versions, requiring careful handling.",
]:
    bullet(con, color=RED)

# ── 4.4 BJS ──────────────────────────────────────────────────────────────────
h2("4.4  Borusyak-Jaravel-Spiess Imputation (BJS) — 04_did_imputation.do")

body(
    "Borusyak, Jaravel, and Spiess (2021, 2024) propose an imputation "
    "estimator. In a first step, the counterfactual outcome under no treatment "
    "is estimated from the pre-treatment observations of all units and the "
    "full panel of never-treated units (using firm and year FEs plus controls). "
    "In a second step, the treatment effect for each treated observation is "
    "estimated as the residual relative to its imputed counterfactual. "
    "Event-study estimates are then formed by averaging these unit-level "
    "estimates within each event-time horizon."
)

h3("Pros")
for pro in [
    "Efficient: by exploiting pre-treatment data of all units for the "
    "counterfactual imputation, it achieves semiparametric efficiency.",
    "Formally targets well-defined ATTs; no negative-weight problem.",
    "Pre-trend test is a joint test of whether the pre-treatment imputation "
    "residuals are zero — a natural and powerful specification test.",
    "Provides heteroskedasticity-robust and cluster-robust inference as "
    "standard.",
]:
    bullet(pro, color=GREEN)

h3("Cons")
for con in [
    "Relies more heavily on the linear factor model for the counterfactual "
    "than non-parametric alternatives like CSDID.",
    "Requires a sufficiently long pre-treatment period for the imputation to "
    "be well-identified.",
    "The Stata package (did_imputation) is less mature than reghdfe and may "
    "have version-specific syntax differences.",
    "Output format varies across package versions; the pipeline therefore logs "
    "ereturn results rather than producing standardised Excel/graph output.",
]:
    bullet(con, color=RED)


# ══════════════════════════════════════════════════════════════════════════════
# 5. DIAGNOSTIC TESTS
# ══════════════════════════════════════════════════════════════════════════════

h1("5. Diagnostic Tests — 05_extra_tests.do")

h2("5.1  Placebo-Shift Pre-Trend Test")
body(
    "The acquisition year G is artificially shifted back by two years for "
    "all treated firms (G_pl = G − 2). The event study is then re-estimated "
    "using the placebo event time reltime_pl = t − G_pl. Under the null of "
    "parallel trends, the pre-treatment coefficients (t_pl = −4, −3, −2) "
    "of this shifted regression should be zero, because those periods correspond "
    "to real pre-treatment years. A significant joint test raises concern about "
    "the parallel-trends assumption."
)

h2("5.2  Bacon Decomposition")
body(
    "Goodman-Bacon (2021) shows that the TWFE DID coefficient is a weighted "
    "average of all pairwise 2×2 DID comparisons: early-treated vs. "
    "never-treated, late-treated vs. never-treated, and (problematically) "
    "late-treated vs. early-treated. The bacondecomp command decomposes the "
    "overall estimate into these components and reports their weights. A large "
    "weight on the 'already-treated vs. later-treated' comparisons signals "
    "that the TWFE estimate may be contaminated by negative-weighting bias."
)

h2("5.3  Wild-Bootstrap Pre-Trend Test")
body(
    "The standard cluster-robust F-test for pre-trends can be unreliable when "
    "the number of clusters is small. The wild cluster bootstrap (boottest, "
    "999 replications, seed 12345) provides an alternative p-value that is "
    "more accurate in small-cluster settings. The joint null is "
    "H₀: d_m4 = d_m3 = d_m2 = 0."
)


# ══════════════════════════════════════════════════════════════════════════════
# 6. COMPARISON FIGURE
# ══════════════════════════════════════════════════════════════════════════════

h1("6. Comparison Figure — 06_comparison_figure.do")
body(
    "The coefficients from TWFE, Sun-Abraham, and CSDID are plotted on the "
    "same axes for the main outcome lnva_L. Each estimator is shown with its "
    "own 95% confidence band. Visual alignment of the three sets of coefficients "
    "provides an informal robustness check: if they diverge substantially in the "
    "post-treatment period this suggests treatment-effect heterogeneity is "
    "economically important and the TWFE estimate is unreliable."
)


# ══════════════════════════════════════════════════════════════════════════════
# 7. HONESTDID SENSITIVITY
# ══════════════════════════════════════════════════════════════════════════════

h1("7. HonestDiD Sensitivity Analysis — 07_honest_did.do")

body(
    "Rambachan and Roth (2023) propose a method to characterise how sensitive "
    "the post-treatment estimates are to violations of the parallel-trends "
    "assumption. Rather than testing whether pre-trends are exactly zero, the "
    "method asks: if the parallel-trends assumption is violated by at most a "
    "certain amount, what is the resulting identified set for the treatment "
    "effect? Two types of restrictions are implemented."
)

h2("7.1  Smoothness Restriction (delta = sd)")
body(
    "The second difference of the counterfactual trend — a measure of how "
    "much the trend accelerates or decelerates — is bounded by M. M = 0 "
    "means exact linear trends (no curvature); larger M allows more deviation. "
    "The pipeline evaluates M ∈ {0, 0.5, 1.0, 1.5, 2.0}. The output plot "
    "shows the identified set for each post-treatment period as a function of M."
)

h2("7.2  Relative Magnitudes (delta = rm)")
body(
    "The maximum post-treatment violation of parallel trends is bounded by "
    "Mbar times the largest pre-treatment deviation observed. Mbar = 0 means "
    "the post-treatment violation cannot exceed the pre-treatment violation; "
    "Mbar = 1 allows it to be equal in magnitude. The pipeline evaluates "
    "Mbar ∈ {0, 0.25, 0.5, 0.75, 1.0}."
)

h2("7.3  Implementation")
body(
    "The TWFE is re-estimated with the same explicit dummy variables as in "
    "01_twfe.do (d_m4, d_m3, d_m2, d_p0 … d_p9). The coefficient vector "
    "e(b)[1..13] and variance matrix e(V)[1..13, 1..13] are passed to honestdid. "
    "The coefplot option produces a sensitivity plot directly. Results are "
    "exported to Excel for further inspection."
)


# ══════════════════════════════════════════════════════════════════════════════
# 8. BUG FIXES
# ══════════════════════════════════════════════════════════════════════════════

h1("8. Bugs Fixed Relative to Original Code")
body(
    "The following bugs were identified during the code review and corrected "
    "in the modular version."
)

tbl4 = doc.add_table(rows=1, cols=4)
tbl4.style = "Table Grid"
h4c = tbl4.rows[0].cells
for i, txt in enumerate(["#", "Location", "Bug", "Fix"]):
    h4c[i].text = txt
    set_cell_bg(h4c[i], "1F3A5F")
    for para in h4c[i].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.bold      = True
            run.font.size      = Pt(10)

bugs = [
    ("1", "All TWFE (01, 05, 07)",
     "ib(-1).reltime##i.treated_firm silently drops all control observations "
     "because reltime = . for controls. Firm and year FEs estimated from "
     "treated units only.",
     "Explicit binary dummies (d_m*, d_p*) set to 0 for controls. Controls "
     "remain in regression throughout."),
    ("2", "02_sun_abraham",
     "control_cohort(.) passes a literal period as a variable name — not "
     "portable across eventstudyinteract versions.",
     "Generate never_treated = (G==.) and pass as control_cohort(never_treated)."),
    ("3", "03_csdid",
     "G = . for controls, but csdid requires G = 0 for never-treated units.",
     "Create G_cs = cond(G<., G, 0) before csdid call."),
    ("4", "03_csdid",
     "local colnames : rownames r(table) returns statistic labels "
     "(b se t p ll ul …), not parameter names.",
     "Changed to colnames r(table)."),
    ("5", "03_csdid",
     "local val = real(regexs(1)) if regexm(…) — Stata's local command "
     "has no if qualifier; regexs(1) evaluated unconditionally.",
     "Replaced with flow-control: if regexm(…) local val = real(regexs(1))."),
    ("6", "04_did_imputation",
     "G = . for controls; did_imputation documentation requires G = 0 "
     "for never-treated.",
     "Create G_bjs = cond(G<., G, 0) before did_imputation call."),
    ("7", "04_did_imputation",
     "pretrends(1/4) passes a numlist; pretrends() takes a single integer.",
     "Changed to pretrends(4)."),
    ("8", "05_extra_tests",
     "capture program drop gen_es_dummies missing; re-running the do-file "
     "errors because the program is already defined.",
     "Added capture program drop gen_es_dummies before program definition."),
    ("9", "Packages (00_prep)",
     "coefplot not installed but required by honestdid, coefplot.",
     "Added coefplot to the package installation loop."),
    ("10", "All modules",
     "sflag_* defined as locals, invisible inside program define blocks "
     "and across do-file boundaries.",
     "Promoted to globals $SFLAG_<outcome> in 00_prep.do."),
]
for num, loc, bug, fix in bugs:
    r = tbl4.add_row().cells
    r[0].text = num
    r[1].text = loc
    r[2].text = bug
    r[3].text = fix
    for para in r[1].paragraphs:
        for run in para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
    set_cell_bg(r[1], "EEF2F7")

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 9. OUTPUTS
# ══════════════════════════════════════════════════════════════════════════════

h1("9. Outputs")

h2("9.1  Excel Files  (Output\\ModernDID\\excel\\)")
for row in [
    ("TWFE_<y>_-4_9.xlsx",       "Columns: tid, b, lb, ub. One file per outcome."),
    ("SA_<y>_-4_9.xlsx",         "Columns: tid, b, lb, ub. One file per outcome."),
    ("CSDID_<y>_-4_9.xlsx",      "Columns: tid, b_cs, lb_cs, ub_cs. One file per outcome."),
    ("HonestDiD_M_<y>.xlsx",     "Smoothness sensitivity table from r(HonestEventStudy)."),
    ("HonestDiD_RM_<y>.xlsx",    "Relative-magnitudes sensitivity table."),
]:
    bullet(row[1], bold_prefix=row[0])

h2("9.2  Graphs  (Output\\ModernDID\\graphs\\)")
for row in [
    ("TWFE_<y>_-4_9.png",        "Event-study plot (navy), 2400 px wide."),
    ("SA_<y>_-4_9.png",          "Event-study plot (dark navy), 2400 px wide."),
    ("CSDID_<y>_-4_9.png",       "Event-study plot (forest green), 2400 px wide."),
    ("COMPARE_lnva_L_-4_9.png",  "Overlay of TWFE + SA + CSDID for lnva_L, 2600 px wide."),
    ("HonestDiD_M_<y>.png",      "Sensitivity plot for smoothness restriction M."),
    ("HonestDiD_RM_<y>.png",     "Sensitivity plot for relative magnitudes Mbar."),
]:
    bullet(row[1], bold_prefix=row[0])

h2("9.3  Logs  (Output\\ModernDID\\logs\\)")
for row in [
    ("sanity_flags.log",             "Tabulations of all seven data-quality checks."),
    ("did_imputation_-4_9.log",      "Full ereturn output from did_imputation for all outcomes."),
]:
    bullet(row[1], bold_prefix=row[0])


# ══════════════════════════════════════════════════════════════════════════════
# 10. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

h1("10. References")

refs = [
    "Borusyak, K., Jaravel, X., & Spiess, J. (2024). Revisiting event-study designs: "
    "Robust and efficient estimation. Review of Economic Studies, rdae007.",

    "Callaway, B., & Sant'Anna, P. H. C. (2021). Difference-in-differences with "
    "multiple time periods. Journal of Econometrics, 225(2), 200–230.",

    "Goodman-Bacon, A. (2021). Difference-in-differences with variation in treatment "
    "timing. Journal of Econometrics, 225(2), 254–277.",

    "Rambachan, A., & Roth, J. (2023). A more credible approach to parallel trends. "
    "Review of Economic Studies, 90(5), 2555–2591.",

    "Sun, L., & Abraham, S. (2021). Estimating dynamic treatment effects in event "
    "studies with heterogeneous treatment effects. Journal of Econometrics, "
    "225(2), 175–199.",
]
for ref in refs:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

# ── Save ─────────────────────────────────────────────────────────────────────
out = "/home/user/fredrikh-economics.github.io/stata/ModernDID_Documentation.docx"
doc.save(out)
print(f"Saved: {out}")
